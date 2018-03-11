
# coding: utf-8

# # 从PDF生成

# In[1]:


if __name__=="__main__":
    input_path='20180309'
    output_path='export'
    
    docx_name='01206765xC4DAF_test.docx'
    txt_name='01207200xC4DAF.txt'
    pdf_name='Report_MHC2397.pdf'
    name_pattern_to_compile=r"(\d+[A-Z]+\d+[A-Z])"


# In[2]:


if __name__=="__main__":
    table_catalog_dict={
        0: {"head": 1, "seqtype":'chain', "chaintype":'HeavyChain'},
        1: {"head": 1, "seqtype":'chain', "chaintype":'LightChain'},
        2: {"head": 1, "seqtype":'CDR', "chaintype":'HC'}, 
        3: {"head": 1, "seqtype":'CDR', "chaintype":'LC'}, 
        4: {"head": 1, "seqtype":'CDR', "chaintype":'HC'},
        5: {"head": 1, "seqtype":'CDR', "chaintype":'LC'},
    }


# ## 载入依赖包
# 
# 在azure notebook上运行时, 由于远程服务器在1小时后会自动关闭并清除所有安装内容, 所以临时安装的依赖软件包并不能正常加载. 需要使用服务器上的预制程序进行设置: 
# 建立script.txt文件: 
# ```txt
# export PATH=~/anaconda3_410/bin:$PATH
# conda install -c anaconda biopython --yes
# pip install python-docx
# ```
# 然后在setting中加载. 这样每次启动azure notebook时会自动安装biopython和python-docx

# In[3]:


# pandas包, 需安装
import pandas as pd
from pandas import DataFrame, Series

# python-docx包, 需安装
from docx import Document

# biopython包, 需安装
from Bio.SeqUtils import seq1



# 系统自带包, 无需安装
from collections import OrderedDict # 有序字典
import re
import os
import numpy as np

# 自建包
from save_sequence import save_seq_to_csv 

# from reconstruct_TXT import reconstruct_TXT
# from reconstruct_DOC import reconstruct_DOC
from reconstruct_PDF import reconstruct_PDF


# In[4]:


if __name__=="__main__":
    fname=pdf_name
    name_pattern=re.compile(name_pattern_to_compile)


# ## 读取文件
# 
# 文件通过两个程序读入: 
# * read_sequence_from_docx 用于读取docx文件, 
# * read_sequence_from_txt 用于读取txt文件
# 
# 由于可能调换source文件和target文件的对应方式, 所以使用了字典source_target_dict来记录对应关系. 
# 使用文件名的末尾3位来判断文件的类型, 是docx/txt
# 
# 文件读取后经过清理, 转换成pandas DataFrame的数据形式进行操作处理. 并且将转换后的DataFrame保存为**同名csv文件**, 存放在output_path文件夹之下

# In[5]:


def get_chains(input_path,pdf_name,output_path,name_pattern):
    pdf_df=reconstruct_PDF(input_path,pdf_name,output_path,name_pattern)
    pdf_df["ID"]=np.arange(len(pdf_df),dtype=np.int32)+1
    HC_df=(pdf_df.where(pdf_df["CHAINTYPE"]=="HeavyChain").dropna())
    LC_df=(pdf_df.where(pdf_df["CHAINTYPE"]=="LightChain").dropna())
    HC_df=(HC_df.drop("CHAINTYPE",axis=1)[["NAME","SEQUENCE","ID"]])
    HC_df.columns=["mAb","VH","ID"]
    HC_df.ID=HC_df.ID.astype("int32")
    LC_df=LC_df.drop("CHAINTYPE",axis=1)[["NAME","SEQUENCE","ID"]]
    LC_df.columns=["mAb","VL","ID"]
    LC_df.ID=LC_df.ID.astype("int32")

    maxID=len(pdf_df)
    return HC_df,LC_df,maxID


# In[6]:


# if __name__=="__main__":
#     fname=pdf_name
#     name_pattern=re.compile(name_pattern_to_compile)
#     HC_df,LC_df,maxID=get_chains(input_path,pdf_name,output_path,name_pattern)


# # 生成CDR

# In[ ]:


def HC_to_CDR(seq):
    '''
    http://www.bioinf.org.uk/abs/#cdrdef
    '''
    CDR1_pattern="(C[A-Z]{3})([A-Z]*?)(W[V|I|A])"
    CDR2_pattern="([A-Z]{12})([A-Z]*?)([K|R])([L|I|V|F|T|A])([T|S|I|A])"
    CDR3_pattern="([A-Z]*)(C[A-Z]{2})([A-Z]{3,25})(WG[A-Z]G)"
    CDRs=re.findall(CDR1_pattern+CDR2_pattern+CDR3_pattern,seq)
#     print(CDRs)
    try:
        CDR1=CDRs[0][1]
        CDR2=CDRs[0][4]
        CDR3=CDRs[0][-2]
    except:
        CDR1=""
        CDR2=""
        CDR3=""
    return [CDR1,CDR2,CDR3]

def LC_to_CDR(seq):
    '''
    假定能够生成CDR
    '''
#     CDR1_pattern="^([A-Z]{20,26}C)([A-Z]{10,17})(W(YQ)|(LQ)|(FQ)|(YL))"
#     CDR2_pattern="([A-Z]{10})((IY)|(VY)|(IK)|(IF))*([A-Z]{7})"
#     CDR3_pattern="([A-Z]{31}C)([A-Z]{7,11})(FG[A-Z]G)"
#     CDRs=re.findall(CDR1_pattern+CDR2_pattern+CDR3_pattern,seq)
    CDR1_pattern="(^[A-Z]{22}C)"+"([A-Z]{10,17})"+"(WYQ|WLQ|WFQ|WYL)"
    CDR2_pattern="([A-Z]*?)"+"([A-Z]{7})"+"(IY|VY|IK|IF)"
    CDR3_pattern="([A-Z]*?C)"+"([A-Z]{7,11})"+"(FG[A-Z]G)"
    CDR_pattern=CDR1_pattern+CDR2_pattern+CDR3_pattern

    CDRs=re.findall(CDR_pattern,seq)
    try:
        CDR1=CDRs[0][1]
        CDR2=CDRs[0][4]
        CDR3=CDRs[0][7]
    except:
        CDR1=""
        CDR2=""
        CDR3=""
    return [CDR1,CDR2,CDR3]


# In[ ]:


seq="DVVMTQTPLSLPVSLGDQASISCRSSQNLVHSYGNTYLHWYLQKPGQSPKLLIYKVSNRFSGVPDRFSGSGSGTDFTLKISRVEAEDLGVYFCSQNTHVPWTFGGGTKLEIQ"
LC_to_CDR(seq)


# In[9]:


def get_CDR_df(select_df,chain_name_base,CDR_method,seq_No_base):
    tempCDRs=select_df.iloc[:,1].apply(CDR_method)
    CDR_index_max=len(tempCDRs)
    CDRs=DataFrame()
    CDRs["mAb"]=select_df["mAb"]

    CDRs["CDR1"]=tempCDRs.apply(lambda x:x[0])
    CDRs["SEQ No. of CDR1"]=np.arange(CDR_index_max)*3+1+seq_No_base

    CDRs["CDR2"]=tempCDRs.apply(lambda x:x[1])
    CDRs["SEQ No. of CDR2"]=np.arange(CDR_index_max)*3+2+seq_No_base

    CDRs["CDR3"]=tempCDRs.apply(lambda x:x[2])
    CDRs["SEQ No. of CDR3"]=np.arange(CDR_index_max)*3+3+seq_No_base

    CDRs.columns=['mAb',
                    chain_name_base+"CDR1","ID",
                    chain_name_base+"CDR2","ID",
                    chain_name_base+"CDR3","ID"]
    maxID=CDR_index_max*3+seq_No_base
    return CDRs,maxID


# In[10]:


# if __name__=="__main__":
#     fname=pdf_name
#     name_pattern=re.compile(name_pattern_to_compile)
    
#     HC_df,LC_df,maxID=get_chains(input_path,pdf_name,output_path,name_pattern)
#     HCCDRs_typeA,maxID=get_CDR_df(HC_df,"HC ",HC_to_CDR,maxID)
#     LCCDRs_typeA,maxID=get_CDR_df(LC_df,"LC ",LC_to_CDR,maxID)
#     HCCDRs_typeB,maxID=get_CDR_df(HC_df,"HC ",HC_to_CDR,maxID)
#     LCCDRs_typeB,maxID=get_CDR_df(LC_df,"LC ",LC_to_CDR,maxID)


# # 生成word文档表格

# In[11]:


def write_table(document,df,table_describe):
    paragraph = document.add_paragraph('\n'+table_describe)
    
    table = document.add_table(rows=df.shape[0]+1, cols=df.shape[1])
#     table.style = 'LightShading-Accent1'
    table.style = 'TableGrid'
    hdr_cells = table.rows[0].cells
    for c,h in zip(df.columns,table.rows[0].cells):
        h.text=c
    for row_idx in range(df.shape[0]):
        row_cells = table.rows[row_idx+1].cells
        for col_idx in range(df.shape[1]):
    #         print("writing ({},{}) with {}".format(row_idx,col_idx,df.iloc[row_idx,col_idx]))
            row_cells[col_idx].text=str(df.iloc[row_idx,col_idx])
    return document


# In[12]:


if __name__=="__main__":
    fname=pdf_name
    name_pattern=re.compile(name_pattern_to_compile)
    
    HC_df,LC_df,maxID=get_chains(input_path,pdf_name,output_path,name_pattern)
    HCCDRs_typeA,maxID=get_CDR_df(HC_df,"HC ",HC_to_CDR,maxID)
    LCCDRs_typeA,maxID=get_CDR_df(LC_df,"LC ",LC_to_CDR,maxID)
    HCCDRs_typeB,maxID=get_CDR_df(HC_df,"HC ",HC_to_CDR,maxID)
    LCCDRs_typeB,maxID=get_CDR_df(LC_df,"LC ",LC_to_CDR,maxID)
    print("all dataframe are done")
    df_list=[HC_df,LC_df,HCCDRs_typeA,LCCDRs_typeA,HCCDRs_typeB,LCCDRs_typeB]
    table_des_list=[
        "Table 1: Sequences of heavy chain variable regions for anti-CD73 mAbs",
        "Table 2: Sequences of light chain variable regions for anti-CD73 mAbs",
        "Table 3: CDR regions 1-3 of heavy chain for anti-CD73 mAbs",
        "Table 4: CDR regions 1-3 of light chain for anti-CD73 mAbs",
        "Table 5: CDR regions 1-3 of heavy chain for anti-CD73 mAbs",
        "Table 6: CDR regions 1-3 of light chain for anti-CD73 mAbs"
    ]
    document = Document()
    for df,table_des in zip(df_list,table_des_list):
        print("write a table")
        document=write_table(document,df,table_des)
    document.save(os.path.join(input_path,output_path,docx_name))


# In[ ]:




